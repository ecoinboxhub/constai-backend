import io
import logging
import time
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Optional

from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from sqlalchemy.exc import SQLAlchemyError

from contextlib import contextmanager
from app.core.config import settings
from app.db.session import SessionLocal
from app.db.models.core import Project, ProjectDocument
from app.modules.project_tracker.schemas import (
    ProjectCreate,
    ProjectResponse,
    ProjectUpdate,
    DashboardMetricsResponse,
    PredictionResponse,
    WeatherResponse,
    DelayPredictionRequest,
    DelayPredictionResponse,
    AIChatRequest,
    AIChatResponse,
    RAGQueryRequest,
    RAGQueryResponse,
    DocumentUploadResponse,
    QuickPredictRequest,
    QuickPredictResponse,
)
from app.services.llm import ask_ai
from app.services.rag.engine import get_vectorstore, query_project_knowledge, chunk_document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
import docx
from app.services.weather import get_live_weather
from app.services.model_service import (
    predict_delay as model_predict_delay,
    predict_budget_overrun as model_predict_budget,
    predict_risk as model_predict_risk,
    extract_project_data,
    extract_quickpredict_data,
    log_prediction,
)

logger = logging.getLogger(__name__)


def _get_redis():
    from app.services.redis_service import get_redis
    return get_redis()


@contextmanager
def get_session():
    session = SessionLocal()
    try:
        yield session
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()


def _safe_float(value: Optional[float], default: float = 0.0) -> float:
    return float(value) if value is not None else default


def _project_to_response(project: Project) -> ProjectResponse:
    return ProjectResponse(
        id=project.id,
        name=project.name,
        contractor_name=project.contractor_name,
        location=project.location,
        state=project.state,
        lga=project.lga,
        project_type=project.project_type,
        start_date=project.start_date,
        expected_end_date=project.expected_end_date,
        actual_end_date=project.actual_end_date,
        project_status=project.project_status,
        budget_allocated=project.budget_allocated or 0.0,
        budget_spent=project.budget_spent,
        workforce_count=project.workforce_count or 0,
        equipment_count=project.equipment_count or 0,
        material_cost=project.material_cost or 0.0,
        completion_percentage=project.completion_percentage or 0.0,
        weather_delay_days=project.weather_delay_days or 0,
        safety_incidents=project.safety_incidents or 0,
        inspection_score=project.inspection_score or 0.0,
        task_completion_rate=project.task_completion_rate or 0.0,
        daily_progress_rate=project.daily_progress_rate or 0.0,
        delay_status=project.delay_status,
        risk_level=project.risk_level,
        company_id=project.company_id,
    )


def list_projects(company_id: int) -> list[ProjectResponse]:
    with get_session() as session:
        projects = session.query(Project).filter(Project.company_id == company_id).order_by(Project.id).all()
        return [_project_to_response(project) for project in projects]


def get_project_by_id(project_id: int, company_id: int) -> Optional[ProjectResponse]:
    with get_session() as session:
        project = session.query(Project).filter(Project.id == project_id, Project.company_id == company_id).first()
        return _project_to_response(project) if project else None


def create_project(payload: ProjectCreate, company_id: int) -> ProjectResponse:
    with get_session() as session:
        try:
            data = payload.model_dump()
            data["company_id"] = company_id
            project = Project(**data)
            session.add(project)
            session.commit()
            session.refresh(project)
            return _project_to_response(project)
        except SQLAlchemyError as exc:
            logger.error(f"Failed to create project: {exc}")
            raise HTTPException(status_code=500, detail="Unable to create project")


def update_project(project_id: int, company_id: int, payload: ProjectUpdate) -> ProjectResponse:
    with get_session() as session:
        try:
            project = session.query(Project).filter(Project.id == project_id, Project.company_id == company_id).first()
            if not project:
                raise HTTPException(status_code=404, detail="Project not found")
            for key, value in payload.model_dump(exclude_unset=True).items():
                setattr(project, key, value)
            session.commit()
            session.refresh(project)
            _get_redis().delete(f"metrics:company:{company_id}")
            return _project_to_response(project)
        except SQLAlchemyError as exc:
            logger.error(f"Failed to update project: {exc}")
            raise HTTPException(status_code=500, detail="Unable to update project")


def delete_project(project_id: int, company_id: int) -> None:
    with get_session() as session:
        project = session.query(Project).filter(Project.id == project_id, Project.company_id == company_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        session.delete(project)
        session.commit()
        _get_redis().delete(f"metrics:company:{company_id}")


def get_dashboard_metrics(company_id: int) -> DashboardMetricsResponse:
    logger.info(f"Fetching dashboard metrics for company {company_id}")
    cache_key = f"metrics:company:{company_id}"
    cached = _get_redis().get(cache_key)
    if cached:
        logger.info(f"Returning cached metrics for company {company_id}")
        return DashboardMetricsResponse(**cached)

    with get_session() as session:
        projects = session.query(Project).filter(Project.company_id == company_id).all()
        total = len(projects)
        if total == 0:
            return DashboardMetricsResponse(
                total_projects=0, active_projects=0, delayed_projects=0,
                completed_projects=0, average_completion=0.0,
                average_budget_utilization=0.0, delay_probability=0.0,
                risk_score=0.0, active_issues=0, productivity_index=0.0,
            )

        completed = sum(1 for p in projects if p.project_status and p.project_status.lower() in {"complete", "completed"})
        active = sum(1 for p in projects if p.project_status and p.project_status.lower() == "active")
        avg_completion = sum((p.completion_percentage or 0.0) for p in projects) / total

        budget_rates = []
        for p in projects:
            if p.budget_allocated and p.budget_allocated > 0:
                rate = (p.budget_spent or 0.0) / p.budget_allocated
                budget_rates.append(min(1.5, rate))

        avg_budget_utilization = sum(budget_rates) / len(budget_rates) if budget_rates else 0.0

        ml_delayed = 0
        total_delay_risk = 0.0
        total_risk_score = 0.0
        total_budget_overrun = 0.0
        ml_issues = 0

        for p in projects:
            try:
                pdata = extract_project_data(p)
                dres = model_predict_delay(pdata)
                pdata["_delay_prob"] = dres["delay_risk"]
                bres = model_predict_budget(pdata)
                rres = model_predict_risk(pdata)

                if dres.get("will_delay", False):
                    ml_delayed += 1

                delay_risk = dres.get("delay_risk", 0.0)
                total_delay_risk += delay_risk

                budget_prob = bres.get("overrun_probability", 0.0)
                total_budget_overrun += budget_prob

                risk_label = rres.get("risk_level", "medium")
                risk_val = {"low": 0.2, "medium": 0.5, "high": 0.85}.get(risk_label.lower(), 0.5)
                total_risk_score += risk_val

                if dres.get("will_delay", False) or (p.safety_incidents or 0) > 0:
                    ml_issues += 1
            except Exception:
                delayed_legacy = p.delay_status and p.delay_status.lower() == "delayed"
                if delayed_legacy:
                    ml_delayed += 1
                    total_delay_risk += 1.0
                else:
                    total_delay_risk += 0.0

                spent = float(getattr(p, 'budget_spent', 0) or 0)
                allocated = float(getattr(p, 'budget_allocated', 1) or 1)
                total_budget_overrun += min(1.0, spent / max(allocated, 1))

                risk_legacy = {"low": 0.2, "medium": 0.5, "high": 0.85}.get((p.risk_level or "medium").lower(), 0.5)
                total_risk_score += risk_legacy

                if delayed_legacy or (p.safety_incidents or 0) > 0:
                    ml_issues += 1

        avg_delay_risk = total_delay_risk / total if total > 0 else 0.0
        avg_risk_score = total_risk_score / total if total > 0 else 0.0

        productivity_index = sum((p.task_completion_rate or (p.completion_percentage / 100 if p.completion_percentage else 0.0)) for p in projects) / total * 100

        res = DashboardMetricsResponse(
            total_projects=total,
            active_projects=active,
            delayed_projects=ml_delayed,
            completed_projects=completed,
            average_completion=round(avg_completion, 2),
            average_budget_utilization=round(avg_budget_utilization * 100, 2),
            delay_probability=round(avg_delay_risk, 3),
            risk_score=round(avg_risk_score * 100, 2),
            active_issues=ml_issues,
            productivity_index=round(productivity_index, 2),
        )
        logger.info(f"Computed fresh metrics for company {company_id}: {res}")
        _get_redis().set(cache_key, res.model_dump(), expire=300)
        return res


def get_weather_for_city(city: str) -> WeatherResponse:
    logger.info(f"Weather request for city: {city}")
    cache_key = f"weather:{city.lower()}"
    cached = _get_redis().get(cache_key)

    if cached:
        logger.info(f"Returning cached weather for {city}")
        return WeatherResponse(city=city.title(), **cached)

    try:
        from app.modules.project_tracker.tasks import fetch_weather_data
        fetch_weather_data.delay(city)
    except Exception as exc:
        logger.warning(f"Background weather refresh unavailable: {exc}")

    logger.info(f"Returning live (non-cached) weather for {city}")
    result = get_live_weather(city)

    if "error" in result:
        logger.warning(f"Weather data unavailable for {city}: {result.get('message', 'Unknown error')}")
        return WeatherResponse(
            city=city.title(),
            temperature_c=0.0,
            condition="Unknown",
            description=result.get("message", "Weather data unavailable"),
            humidity_pct=0,
            pressure_hpa=0,
            wind_speed_kmh=0,
            rainfall_mm=0,
            fetched_at=time.time(),
            source="ConstAI (Data Unavailable)"
        )

    return WeatherResponse(city=city.title(), **result)


def predict_delay(payload: DelayPredictionRequest) -> DelayPredictionResponse:
    live_weather = None
    if payload.city:
        live_weather = get_live_weather(payload.city)
        if payload.rainfall_mm == 0 and live_weather.get("rainfall_mm", 0) > 0:
            payload.rainfall_mm = live_weather["rainfall_mm"]
        if payload.temperature_c is None:
            payload.temperature_c = live_weather.get("temperature_c")
        if payload.wind_speed_kmh is None:
            payload.wind_speed_kmh = live_weather.get("wind_speed_kmh")

    score = 0.05
    score += min(0.35, payload.rainfall_mm * 0.02)
    score += min(0.25, (1.0 - payload.resource_availability) * 0.35)
    score += min(0.25, (1.0 - payload.workforce_attendance) * 0.35)
    score += min(0.20, payload.supply_delay_days * 0.05)

    if payload.temperature_c is not None:
        score += 0.05 if payload.temperature_c > 32 else 0.0
    if payload.wind_speed_kmh is not None:
        score += 0.03 if payload.wind_speed_kmh > 25 else 0.0

    if live_weather and live_weather.get("severe_alert"):
        score = min(1.0, score + 0.20)

    delay_risk = min(1.0, round(score, 4))

    if delay_risk > 0.75:
        advisory = "High delay risk: accelerate supply chain and increase crew coverage."
        if live_weather and live_weather.get("severe_alert"):
            advisory += f" {live_weather['severe_alert']}"
    elif delay_risk > 0.4:
        advisory = "Moderate delay risk: monitor weather and resource availability closely."
    else:
        advisory = "Low delay risk: continue current schedule and verify deliveries."

    estimated_delay_days = int(round(payload.supply_delay_days * (0.8 + delay_risk)))
    return DelayPredictionResponse(
        delay_risk=delay_risk,
        advisory=advisory,
        estimated_delay_days=max(0, estimated_delay_days),
    )


def _heuristic_budget_probability(project: Project) -> float:
    spent = _safe_float(project.budget_spent)
    allocated = _safe_float(project.budget_allocated, 1.0)
    trend = min(1.0, spent / max(allocated, 1.0))
    score = 0.2 + 0.4 * trend
    score += 0.15 if (project.weather_delay_days or 0) > 5 else 0.0
    score += 0.15 if (project.safety_incidents or 0) > 2 else 0.0
    return min(1.0, score)


def _classify_risk(project: Project, delay_prob: float = 0.0) -> str:
    if project.risk_level:
        return project.risk_level

    budget_prob = _heuristic_budget_probability(project)
    score = delay_prob + budget_prob

    if score > 1.2:
        return "high"
    if score > 0.7:
        return "medium"
    return "low"


async def get_project_predictions(project_id: int, company_id: int) -> PredictionResponse:
    logger.info(f"Fetching AI intelligence for project {project_id}, company {company_id}")
    with get_session() as session:
        project = session.query(Project).filter(Project.id == project_id, Project.company_id == company_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        weather_context = get_live_weather(project.location or "Lagos")
        if "error" in weather_context:
            logger.warning(f"Live weather unavailable for {project.location}: {weather_context['message']}")

        project_data = extract_project_data(project)

        delay_result = model_predict_delay(project_data)
        delay_prob = delay_result["delay_risk"]

        project_data["_delay_prob"] = delay_prob
        budget_result = model_predict_budget(project_data)
        budget_prob = budget_result["overrun_probability"]

        risk_result = model_predict_risk(project_data)
        risk_label = risk_result["risk_level"]

        log_prediction(project_data, delay_result, project_id=project.id, company_id=company_id)

        completion_forecast = min(100.0, _safe_float(project.completion_percentage) + _safe_float(project.daily_progress_rate) * 7.0)
        cost_trend = min(2.0, (_safe_float(project.budget_spent) / max(_safe_float(project.budget_allocated), 1.0)))
        estimated_date = project.actual_end_date or project.expected_end_date

        return PredictionResponse(
            project_id=project.id,
            delay_probability=round(delay_prob, 4),
            budget_overrun_probability=round(budget_prob, 4),
            risk_classification=risk_label,
            delay_model_version=delay_result.get("model_version", "v0"),
            budget_model_version=budget_result.get("model_version", "v0"),
            risk_model_version=risk_result.get("model_version", "v0"),
            estimated_completion_date=estimated_date,
            completion_forecast=round(completion_forecast, 2),
            cost_trend=round(cost_trend, 3),
            weather_context=WeatherResponse(city=project.location or "Lagos", **weather_context) if "error" not in weather_context else None
        )


async def chat_insight(payload: AIChatRequest, company_id: int) -> AIChatResponse:
    project_context_data = {"company_id": company_id, "timestamp": datetime.now(UTC).isoformat()}

    if payload.project_id:
        project = get_project_by_id(payload.project_id, company_id)
        if not project:
            projects = list_projects(company_id)
            summary = f"Portfolio for Org {company_id} has {len(projects)} projects."
            project_context_data["project_count"] = len(projects)
            payload.project_id = None
        else:
            summary = f"Project '{project.name}' (ID: {project.id}) at {project.location}. Status: {project.project_status}, Completion: {project.completion_percentage}%, Budget: NGN{project.budget_allocated:,.0f}"
        project_context_data["project_id"] = payload.project_id
        project_context_data["project_name"] = project.name
    else:
        projects = list_projects(company_id)
        summary = f"Portfolio for Org {company_id} has {len(projects)} projects."
        project_context_data["project_count"] = len(projects)

    prompt = f"""You are a Construction Project Intelligence Assistant for ConstAI Nigeria.

{summary}

User Query: {payload.message}

IMPORTANT INSTRUCTIONS:
- Provide your response as plain natural language text only
- Do NOT use markdown formatting (no **bold**, no *italics*, no # headers, no - bullets)
- Do NOT use JSON syntax or any code formatting
- Do NOT use numbered lists or bullet points
- Write in simple, conversational paragraphs
- Just provide clear, straightforward text

Response:"""
    response = await ask_ai(prompt)

    return AIChatResponse(
        response=response,
        project_context=project_context_data
    )


async def rag_query(payload: RAGQueryRequest, company_id: int) -> RAGQueryResponse:
    persist_dir = Path(settings.chroma_persist_dir) / "project_documents"
    result = await query_project_knowledge(
        project_id=payload.project_id,
        company_id=company_id,
        question=payload.question,
        persist_dir=str(persist_dir)
    )

    return RAGQueryResponse(
        answer=result["answer"],
        sources=result.get("sources", []),
        source_count=len(result.get("sources", [])),
        timestamp=datetime.now(UTC).isoformat()
    )


def _extract_text(file_name: str, raw_bytes: bytes) -> str:
    lower_name = file_name.lower()
    if lower_name.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(raw_bytes))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            pass
    if lower_name.endswith(".docx"):
        try:
            document = docx.Document(io.BytesIO(raw_bytes))
            return "\n\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            pass
    return raw_bytes.decode("utf-8", errors="ignore")


async def upload_project_document(project_id: int, company_id: int, file: UploadFile) -> DocumentUploadResponse:
    with get_session() as session:
        project = session.query(Project).filter(Project.id == project_id, Project.company_id == company_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        raw_bytes = await file.read()
        text = _extract_text(file.filename, raw_bytes)

        storage_dir = Path("data/processed/rag_documents")
        storage_dir.mkdir(parents=True, exist_ok=True)
        file_path = storage_dir / f"{company_id}_{project_id}_{file.filename}"
        file_path.write_bytes(raw_bytes)

        document = ProjectDocument(
            project_id=project_id,
            company_id=company_id,
            file_name=file.filename,
            source_path=str(file_path),
            content=text,
        )
        session.add(document)
        session.commit()
        session.refresh(document)

        chunks = chunk_document(text)
        persist_dir = Path(settings.chroma_persist_dir) / "project_documents"
        persist_dir.mkdir(parents=True, exist_ok=True)
        vectorstore = get_vectorstore(str(persist_dir))

        vectorstore.add_texts(
            texts=chunks,
            metadatas=[{
                "source": file.filename,
                "project_id": project_id,
                "company_id": company_id
            }] * len(chunks)
        )

        return DocumentUploadResponse(
            project_id=project_id,
            file_name=file.filename,
            indexed_chunks=len(chunks),
            message="Document uploaded and indexed successfully"
        )


async def quick_predict(payload: QuickPredictRequest) -> QuickPredictResponse:
    data = extract_quickpredict_data(payload)

    delay_result = model_predict_delay(data)
    delay_prob = delay_result["delay_risk"]

    budget_result = model_predict_budget(data)
    budget_prob = budget_result["overrun_probability"]

    data["_delay_prob"] = delay_prob
    risk_result = model_predict_risk(data)
    risk_label = risk_result["risk_level"]

    key_factors = []
    if delay_prob > 0.5:
        key_factors.append("High delay probability")
    if budget_prob > 0.5:
        key_factors.append("High budget overrun risk")
    if payload.weather_delay_days > 10:
        key_factors.append(f"Weather delays ({payload.weather_delay_days} days)")
    if payload.safety_incidents > 3:
        key_factors.append(f"Safety incidents ({payload.safety_incidents})")
    if payload.task_completion_rate < 0.5:
        key_factors.append("Low task completion rate")
    if not key_factors:
        key_factors.append("Project appears on track")

    advisory_prompt = (
        f"Project: {payload.completion_percentage}% complete, "
        f"delay probability {delay_prob:.0%}, budget overrun probability {budget_prob:.0%}, "
        f"risk level {risk_label}. "
        f"Factors: {', '.join(key_factors)}. "
        "Provide concise construction advisory (2-3 sentences)."
    )
    advisory = await ask_ai(advisory_prompt)

    return QuickPredictResponse(
        delay_probability=round(delay_prob, 4),
        budget_overrun_probability=round(budget_prob, 4),
        risk_level=risk_label,
        model_version=delay_result.get("model_version", "v0"),
        advisory=advisory,
        key_risk_factors=key_factors,
    )
